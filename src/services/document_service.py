"""
Servicio de gestión de documentos y biblioteca
"""

import os
import uuid
import aiofiles
from pathlib import Path
from typing import List, Optional, Dict, Any
import PyPDF2
import docx
import pandas as pd
from datetime import datetime

from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings as ChromaSettings

from src.config import settings
from src.models import Document, DocumentType


class DocumentService:
    """Servicio para gestión de documentos y búsqueda semántica"""
    
    def __init__(self):
        self.embedding_model = SentenceTransformer(settings.embedding_model)
        
        # Configurar ChromaDB
        self.chroma_client = chromadb.PersistentClient(
            path=settings.chroma_persist_directory,
            settings=ChromaSettings(anonymized_telemetry=False)
        )
        
        # Crear colección si no existe
        try:
            self.collection = self.chroma_client.get_collection("documents")
        except:
            self.collection = self.chroma_client.create_collection(
                name="documents",
                metadata={"description": "Documentos educativos"}
            )
    
    async def upload_document(self, file_content: bytes, filename: str, 
                            subject: Optional[str] = None, 
                            grade_level: Optional[str] = None) -> Document:
        """Subir y procesar un documento"""
        
        # Generar ID único
        doc_id = str(uuid.uuid4())
        
        # Determinar tipo de archivo
        file_extension = filename.split('.')[-1].lower()
        try:
            file_type = DocumentType(file_extension)
        except ValueError:
            raise ValueError(f"Tipo de archivo no soportado: {file_extension}")
        
        # Guardar archivo físico
        file_path = settings.upload_dir / f"{doc_id}_{filename}"
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Extraer contenido del texto
        content = await self._extract_text(file_path, file_type)
        
        # Crear documento
        document = Document(
            id=doc_id,
            filename=filename,
            file_type=file_type,
            content=content,
            subject=subject,
            grade_level=grade_level,
            metadata={
                "file_path": str(file_path),
                "file_size": len(file_content)
            }
        )
        
        # Procesar e indexar
        await self._index_document(document)
        
        return document
    
    async def _extract_text(self, file_path: Path, file_type: DocumentType) -> str:
        """Extraer texto de diferentes tipos de archivo"""
        
        if file_type == DocumentType.TXT or file_type == DocumentType.MARKDOWN:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        
        elif file_type == DocumentType.PDF:
            return self._extract_pdf_text(file_path)
        
        elif file_type == DocumentType.DOCX:
            return self._extract_docx_text(file_path)
        
        elif file_type == DocumentType.XLSX:
            return self._extract_xlsx_text(file_path)
        
        else:
            raise ValueError(f"Tipo de archivo no soportado: {file_type}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extraer texto de PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extrayendo texto de PDF: {e}")
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extraer texto de DOCX"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extrayendo texto de DOCX: {e}")
        return text
    
    def _extract_xlsx_text(self, file_path: Path) -> str:
        """Extraer texto de XLSX"""
        text = ""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            for sheet_name, sheet_df in df.items():
                text += f"Hoja: {sheet_name}\n"
                text += sheet_df.to_string() + "\n\n"
        except Exception as e:
            print(f"Error extrayendo texto de XLSX: {e}")
        return text
    
    async def _index_document(self, document: Document):
        """Indexar documento en la base de datos vectorial"""
        
        # Dividir el contenido en chunks
        chunks = self._split_text(document.content)
        
        # Generar embeddings
        embeddings = self.embedding_model.encode(chunks)
        
        # Preparar metadatos
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{document.id}_chunk_{i}"
            ids.append(chunk_id)
            
            metadata = {
                "document_id": document.id,
                "filename": document.filename,
                "file_type": document.file_type.value,
                "chunk_index": i,
                "subject": document.subject or "",
                "grade_level": document.grade_level or "",
                "uploaded_at": document.uploaded_at.isoformat()
            }
            metadatas.append(metadata)
        
        # Añadir a ChromaDB
        self.collection.add(
            embeddings=embeddings.tolist(),
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
    
    def _split_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Dividir texto en chunks con overlap"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap
            
            if start >= len(text):
                break
        
        return chunks
    
    async def search_documents(self, query: str, n_results: int = 10, 
                             subject: Optional[str] = None,
                             grade_level: Optional[str] = None) -> List[Dict[str, Any]]:
        """Buscar documentos usando búsqueda semántica"""
        
        # Generar embedding de la consulta
        query_embedding = self.embedding_model.encode([query])
        
        # Preparar filtros
        where_filter = {}
        if subject:
            where_filter["subject"] = subject
        if grade_level:
            where_filter["grade_level"] = grade_level
        
        # Buscar en ChromaDB
        results = self.collection.query(
            query_embeddings=query_embedding.tolist(),
            n_results=n_results,
            where=where_filter if where_filter else None
        )
        
        # Formatear resultados
        formatted_results = []
        for i in range(len(results['ids'][0])):
            result = {
                "id": results['ids'][0][i],
                "document": results['documents'][0][i],
                "metadata": results['metadatas'][0][i],
                "distance": results['distances'][0][i] if 'distances' in results else None
            }
            formatted_results.append(result)
        
        return formatted_results
    
    async def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """Obtener documento por ID"""
        
        results = self.collection.get(
            where={"document_id": document_id},
            limit=1
        )
        
        if not results['ids']:
            return None
        
        metadata = results['metadatas'][0]
        
        # Reconstruir documento
        document = Document(
            id=document_id,
            filename=metadata['filename'],
            file_type=DocumentType(metadata['file_type']),
            content="",  # El contenido completo se obtendría del archivo
            subject=metadata.get('subject'),
            grade_level=metadata.get('grade_level'),
            uploaded_at=datetime.fromisoformat(metadata['uploaded_at'])
        )
        
        return document
    
    async def list_documents(self, subject: Optional[str] = None, 
                           grade_level: Optional[str] = None) -> List[Dict[str, Any]]:
        """Listar todos los documentos"""
        
        where_filter = {}
        if subject:
            where_filter["subject"] = subject
        if grade_level:
            where_filter["grade_level"] = grade_level
        
        results = self.collection.get(
            where=where_filter if where_filter else None
        )
        
        # Agrupar por document_id para evitar duplicados
        documents = {}
        for i, doc_id in enumerate(results['ids']):
            metadata = results['metadatas'][i]
            document_id = metadata['document_id']
            
            if document_id not in documents:
                documents[document_id] = {
                    "id": document_id,
                    "filename": metadata['filename'],
                    "file_type": metadata['file_type'],
                    "subject": metadata.get('subject', ''),
                    "grade_level": metadata.get('grade_level', ''),
                    "uploaded_at": metadata['uploaded_at']
                }
        
        return list(documents.values())
    
    async def delete_document(self, document_id: str) -> bool:
        """Eliminar documento"""
        
        try:
            # Eliminar de ChromaDB
            self.collection.delete(where={"document_id": document_id})
            
            # Eliminar archivo físico
            results = self.collection.get(
                where={"document_id": document_id},
                limit=1
            )
            
            if results['metadatas']:
                metadata = results['metadatas'][0]
                file_path = Path(metadata.get('file_path', ''))
                if file_path.exists():
                    file_path.unlink()
            
            return True
        
        except Exception as e:
            print(f"Error eliminando documento: {e}")
            return False
