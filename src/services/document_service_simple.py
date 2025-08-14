"""
Servicio simplificado de gestión de documentos
"""

import os
import uuid
import aiofiles
from pathlib import Path
from typing import List, Optional, Dict, Any
import PyPDF2
import docx
from datetime import datetime

from src.config import settings
from src.models import Document, DocumentType


class DocumentService:
    """Servicio simplificado de gestión de documentos sin ChromaDB"""
    
    def __init__(self):
        # Crear directorios necesarios
        settings.upload_dir.mkdir(parents=True, exist_ok=True)
        # Almacenamiento simple en memoria (en producción usar base de datos)
        self.documents = []
    
    async def upload_document(self, file_content: bytes, filename: str, 
                            subject: str = None, grade_level: str = None) -> Document:
        """Subir un documento de manera simplificada"""
        
        # Generar ID único
        doc_id = str(uuid.uuid4())
        
        # Determinar tipo de documento
        file_extension = filename.split('.')[-1].lower()
        doc_type = DocumentType(file_extension) if file_extension in [t.value for t in DocumentType] else DocumentType.TXT
        
        # Guardar archivo
        file_path = settings.upload_dir / f"{doc_id}_{filename}"
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Extraer contenido del texto
        content = await self._extract_text_content(file_path, doc_type)
        
        # Crear documento
        document = Document(
            id=doc_id,
            filename=filename,
            file_path=str(file_path),
            document_type=doc_type,
            content=content,
            subject=subject,
            grade_level=grade_level,
            upload_date=datetime.now(),
            metadata={
                "file_size": len(file_content),
                "original_filename": filename
            }
        )
        
        # Almacenar en memoria (temporal)
        self.documents.append(document)
        
        return document
    
    async def _extract_text_content(self, file_path: Path, doc_type: DocumentType) -> str:
        """Extraer contenido de texto del archivo"""
        
        try:
            if doc_type == DocumentType.PDF:
                return await self._extract_pdf_text(file_path)
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx_text(file_path)
            elif doc_type == DocumentType.TXT:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    return await f.read()
            else:
                return "Contenido no extraíble para este tipo de archivo"
        except Exception as e:
            return f"Error extrayendo contenido: {str(e)}"
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extraer texto de PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Error leyendo PDF: {str(e)}"
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extraer texto de DOCX"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error leyendo DOCX: {str(e)}"
    
    async def list_documents(self, subject: str = None, grade_level: str = None) -> List[Document]:
        """Listar documentos con filtros opcionales"""
        
        filtered_docs = self.documents
        
        if subject:
            filtered_docs = [doc for doc in filtered_docs if doc.subject == subject]
        
        if grade_level:
            filtered_docs = [doc for doc in filtered_docs if doc.grade_level == grade_level]
        
        return filtered_docs
    
    async def search_documents(self, query: str, n_results: int = 10, 
                             subject: str = None, grade_level: str = None) -> List[Dict[str, Any]]:
        """Búsqueda simple por texto (sin embeddings)"""
        
        results = []
        query_lower = query.lower()
        
        for doc in self.documents:
            # Filtrar por subject y grade_level si se especifican
            if subject and doc.subject != subject:
                continue
            if grade_level and doc.grade_level != grade_level:
                continue
            
            # Búsqueda simple por contenido
            if query_lower in doc.content.lower():
                results.append({
                    "id": doc.id,
                    "filename": doc.filename,
                    "content": doc.content,
                    "subject": doc.subject,
                    "grade_level": doc.grade_level,
                    "metadata": doc.metadata,
                    "relevance_score": doc.content.lower().count(query_lower)  # Score simple
                })
        
        # Ordenar por relevancia (frecuencia de aparición)
        results.sort(key=lambda x: x["relevance_score"], reverse=True)
        
        return results[:n_results]
    
    async def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """Obtener documento por ID"""
        for doc in self.documents:
            if doc.id == document_id:
                return doc
        return None
    
    async def get_documents_by_ids(self, document_ids: List[str]) -> List[Document]:
        """Obtener múltiples documentos por IDs"""
        return [doc for doc in self.documents if doc.id in document_ids]
    
    async def delete_document(self, document_id: str) -> bool:
        """Eliminar documento"""
        
        # Buscar documento
        doc_to_delete = None
        for i, doc in enumerate(self.documents):
            if doc.id == document_id:
                doc_to_delete = doc
                del self.documents[i]
                break
        
        if doc_to_delete:
            # Eliminar archivo físico
            try:
                file_path = Path(doc_to_delete.file_path)
                if file_path.exists():
                    file_path.unlink()
                return True
            except Exception as e:
                print(f"Error eliminando archivo: {e}")
                return False
        
        return False
    
    def get_library_stats(self) -> Dict[str, Any]:
        """Obtener estadísticas de la biblioteca"""
        
        subjects = {}
        grade_levels = {}
        doc_types = {}
        
        for doc in self.documents:
            # Contar por materia
            if doc.subject:
                subjects[doc.subject] = subjects.get(doc.subject, 0) + 1
            
            # Contar por nivel
            if doc.grade_level:
                grade_levels[doc.grade_level] = grade_levels.get(doc.grade_level, 0) + 1
            
            # Contar por tipo
            doc_types[doc.document_type.value] = doc_types.get(doc.document_type.value, 0) + 1
        
        return {
            "total_documents": len(self.documents),
            "by_subject": subjects,
            "by_grade_level": grade_levels,
            "by_document_type": doc_types,
            "total_size_mb": sum([doc.metadata.get("file_size", 0) for doc in self.documents]) / (1024*1024)
        }
